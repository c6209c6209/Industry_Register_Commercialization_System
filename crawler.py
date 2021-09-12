import urllib.request as req
from bs4 import BeautifulSoup
from os import path, mkdir
from xlsxwriter import Workbook
from docx import Document
from docx.shared import Cm, Mm, Pt
from docx.table import Table
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
import zipcodetw

# PROJECT_PATH = path.realpath(path.dirname(__file__))
# ANYFILE = path.join(PROJECT_PATH, 'db/for_example_db')

cities = {"基隆市" : "10017",
        "宜蘭縣" : "10002",
        "臺北市" : "63",
        "新北市" : "10001",
        "桃園市" : "10003",
        "新竹縣" : "10004",
        "新竹市" : "10018",
        "苗栗縣" : "10005",
        "臺中市" : "10019",
        "彰化縣" : "10007",
        "南投縣" : "10008",
        "雲林縣" : "10009",
        "嘉義縣" : "10010",
        "嘉義市" : "10020",
        "臺南市" : "10021",
        "高雄市" : "64",
        "屏東縣" : "10013",
        "花蓮縣" : "10015",
        "臺東縣" : "10014",
        "澎湖縣" : "10016",
        "金門縣" : "09020",
        "連江縣" : "09007"}

def getAgency(url):
    # get the page
    headers = {"User-Agent" : "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.116 Safari/537.36"}
    request = req.Request(url, headers = headers)
    with req.urlopen(request) as response:
        data = response.read().decode("utf-8")

    # get name and address of each agency
    root = BeautifulSoup(data, "html.parser")
    # agencies = root.find_all("ul", class_="subhotelList")
    agencies = root.find_all("li", {"aria-label" : ["名稱" , "地址", "電話"]})
    
    for agency in agencies:
        try:
            if agency["aria-label"] == "名稱":
                cur_agency = {}
                cur_agency["name"] = agency.string
            elif agency["aria-label"] == "地址" :
                # Check if it has multiple floors. If yes, only get the first floor
                multiple_floor_index = agency.string.find("、")
                if multiple_floor_index == -1:
                    cur_agency["address"] = agency.string
                else :
                    cur_agency["address"] = agency.string[0:multiple_floor_index]
                    if agency.string[multiple_floor_index - 1] != "樓" and agency.string[multiple_floor_index - 1] != "號" and agency.string[multiple_floor_index - 2] != "之" and agency.string[multiple_floor_index + 1] != "樓" :
                        cur_agency["address"] = cur_agency["address"] + "樓"
                # remove parentheses
                parentheses = "()（）﹝﹞"
                for char in parentheses:   
                    cur_agency["address"] = cur_agency["address"].replace(char,"")
                #find zip code
                cur_agency["zipcode"] = zipcodetw.find(agency.string)
            elif agency["aria-label"] == "電話" :
                cur_agency["telephone"] = agency.string
                # add the travel agency to list
                agency_list.append(cur_agency)
        except KeyError:
            pass
    
    # get url of next page
    next_page = root.find("a", class_ = "next-page")
    if next_page != None:
        return next_page['href']
    else:  # the last page
        return ""

# create a directory to put data of agencies in xlsx file
dirName = '旅行社名單'
try:
    mkdir(dirName)
    print("Directory " + dirName + " created.\n") 
except FileExistsError:
    print("Directory " + dirName + " already exists.\n")
# create a directory to put mail typesetting to docx file
dirName = '旅行社郵寄排版'
try:
    mkdir(dirName)
    print("Directory " + dirName + " created.\n") 
except FileExistsError:
    print("Directory " + dirName + " already exists.\n")


for city, value in cities.items():
    print("Catching " + city + " data ...")
    agency_list = []
    url = "https://www.taiwan.net.tw/m1.aspx?sNo=0000148&lid=53&keyString=%5e"+ value +"%5e%5e%5e1"
    while url != "https://www.taiwan.net.tw/":
        url = "https://www.taiwan.net.tw/"+getAgency(url)

    # write data of agencies in xlsx file
    print("Writing data to " + city + " xlsx file ...")
    wb = Workbook("旅行社名單\\" + city + ".xlsx")
    sheet = wb.add_worksheet("Sheet 1")
    sheet.set_column(0, 0, 31.9)
    sheet.set_column(1, 1, 9.7)
    sheet.set_column(2, 2, 45)
    sheet.set_column(3, 3, 13.1)
    title_style = wb.add_format({"bold": True})
    sheet.write(0, 0, "公司名稱", title_style)
    sheet.write(0, 1, "郵遞區號", title_style)
    sheet.write(0, 2, "地址", title_style)
    sheet.write(0, 3, "電話", title_style)
    for row, agency in enumerate(agency_list):
        sheet.write(row+1, 0, agency["name"])
        sheet.write(row+1, 1, agency["zipcode"])
        sheet.write(row+1, 2, agency["address"])
        sheet.write(row+1, 3, agency["telephone"])
    wb.close()

    # write mail typesetting to docx file
    print("Writing data to " + city + " docx file ...")
    doc = Document()
    # set page margin
    for section in doc.sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0.01)
        section.page_height = Mm(297)
        section.page_width = Mm(210)
    # fill agencies in each table entry
    for num, agency in enumerate(agency_list):
        if num % 36 == 0:
            # create table
            table = doc.add_table(12, 3)
            for col in table.columns:
                col.width = Mm(70.8)
            table.rows[0].height = Mm(24)
            for row in table.rows[1:11]:
                row.height = Mm(24.8)
            table.rows[11].height = Mm(24)
            # set line spacing to 1pt
            paragraph = doc.add_paragraph("")
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = Pt(1)
        # edit table
        row = int((num%36)/3)
        col = int((num%36)%3)
        if (num%36)%3 == 0:
            table.cell(row, col).text = "  " + agency["name"] + "\n  (" + agency["zipcode"] + ")" + agency["address"] + "\n  採購人員  收"
            # table.cell(row, col).text = "  " + agency["name"] + "  收\n  "  + agency["address"]
        else:
            table.cell(row, col).text = agency["name"] + "\n(" + agency["zipcode"] + ")" + agency["address"] + "\n採購人員  收"
            # table.cell(row, col).text = agency["name"] + "  收\n"  + agency["address"]
        table.cell(row,col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in table.cell(row, col).paragraphs:
            for run in paragraph.runs:
                run.font.name = "標楷體"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
                address_start = run.text.find("社")
                if len(run.text) - address_start < 28:
                    run.font.size = Pt(8)
                else:
                    run.font.size = Pt(7)

    # save the file
    doc.save("旅行社郵寄排版\\" + city + ".docx")

    print("Finished writing " + city + " data !")
    print("\n")